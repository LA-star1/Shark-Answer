"""FastAPI application — main entry point for Shark Answer."""

from __future__ import annotations

import io
import json
import logging
import re
import time
import uuid
import zipfile
from contextlib import asynccontextmanager
from dataclasses import asdict
from pathlib import Path
from typing import AsyncGenerator, Optional

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel

from shark_answer.config import (
    AppConfig, Language, Pipeline, Subject, SUBJECT_PIPELINE_MAP,
)
from shark_answer.knowledge_base.retriever import build_prompt_context
from shark_answer.modules.examiner_profile import ExaminerProfileManager
from shark_answer.pipelines.base import PipelineResult
from shark_answer.pipelines.pipeline_a_science import run_pipeline_a
from shark_answer.pipelines.pipeline_b_essay import run_pipeline_b
from shark_answer.pipelines.pipeline_c_cs import run_pipeline_c
from shark_answer.pipelines.pipeline_d_practical import run_pipeline_d
from shark_answer.pipelines.router import route_question
from shark_answer.providers.registry import ProviderRegistry
from shark_answer.utils.cost_tracker import CostTracker
from shark_answer.utils.file_converter import convert_file_to_images
from shark_answer.utils.image_extractor import (
    extract_questions_from_images,
    ExtractedQuestion,
)

logger = logging.getLogger(__name__)

# Global state
_config: Optional[AppConfig] = None
_registry: Optional[ProviderRegistry] = None
_examiner_manager: Optional[ExaminerProfileManager] = None
_cost_tracker: Optional[CostTracker] = None

# In-memory history store (list of past submissions)
_history: list[dict] = []

# Chat history: key = "submission_id:question_number", value = list of messages
_chat_histories: dict[str, list[dict]] = {}

# Last-run log for /api/debug/last-run
_last_run_log: dict = {}

# Template and static paths
_BASE_DIR = Path(__file__).parent
_TEMPLATES_DIR = _BASE_DIR / "templates"
_STATIC_DIR = _BASE_DIR / "static"


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize application state on startup."""
    global _config, _registry, _examiner_manager, _cost_tracker

    _config = AppConfig.from_env()
    logging.basicConfig(level=getattr(logging, _config.log_level))

    _registry = ProviderRegistry(_config)
    _examiner_manager = ExaminerProfileManager(_config.examiner_profile_dir)
    _cost_tracker = CostTracker(budget_warning_usd=_config.cost_budget_warning_usd)

    available = list(_config.models.keys())
    logger.info("Shark Answer started. Available models: %s",
                [m.value for m in available])

    yield

    logger.info("Shark Answer shutting down. Total cost: $%.4f",
                _cost_tracker.total_cost if _cost_tracker else 0)


app = FastAPI(
    title="Shark Answer",
    description="A-Level CIE Exam Answer Generation System",
    version="0.1.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory=str(_STATIC_DIR)), name="static")

# Templates
templates = Jinja2Templates(directory=str(_TEMPLATES_DIR))


# ===== Pydantic models =====

class AnswerVersionOut(BaseModel):
    version_number: int
    answer_text: str
    explanation_text: str
    approach_label: str
    provider: str
    verified: bool
    quality_score: Optional[str] = None   # e.g. "14/15"
    language: str


class QuestionResultOut(BaseModel):
    question_number: str
    question_text: str
    pipeline: str
    subject: str
    versions: list[AnswerVersionOut]
    verification_notes: str
    disagreement_resolved: bool
    errors: list[str]


class PaperResultOut(BaseModel):
    total_questions: int
    results: list[QuestionResultOut]
    cost_summary: dict


class ExaminerProfileOut(BaseModel):
    name: str
    subject: str
    region: str
    description: str


class CostSummaryOut(BaseModel):
    total_cost_usd: float
    total_calls: int
    total_input_tokens: int
    total_output_tokens: int
    by_provider: dict[str, float]
    by_subject: dict[str, float]


class ChatRequest(BaseModel):
    submission_id: str
    question_number: str
    message: str


class ChatResponse(BaseModel):
    reply: str
    chat_history: list[dict]


# ===========================================================
# Scoring & revision prompts
# ===========================================================

SCORE_SYSTEM = """You are a CIE A-Level examiner scoring a student answer against the mark scheme.

Output ONLY valid JSON with this exact structure (no markdown fences):
{
  "marks_achieved": <integer>,
  "total_marks": <integer>,
  "verdict": "full" | "partial" | "poor",
  "missing_points": ["specific mark point not yet addressed", ...]
}

Rules:
- verdict is "full"    if marks_achieved == total_marks
- verdict is "partial" if marks_achieved >= ceil(total_marks * 0.7)
- verdict is "poor"    otherwise
- missing_points lists specific CIE mark-scheme criteria NOT yet addressed (empty list when verdict=="full")"""

REVISE_SYSTEM = """You are a top CIE A-Level student improving your exam answer to earn the missing marks indicated by the examiner.

Rewrite the COMPLETE answer incorporating the missing points. Keep the same style and approximate length.
Write ONLY the improved answer — no preamble, no "Here is the revised..." header, no commentary."""


async def _score_and_revise(
    question_text: str,
    marks: int,
    answer_text: str,
    subject: str,
    language: str = "en",
) -> tuple[str, str]:
    """Score an answer with Claude and revise up to 3 rounds.

    Returns (best_answer_text, score_str) where score_str is e.g. "14/15".
    Falls back gracefully if Claude is unavailable.
    """
    if not _registry or not _cost_tracker:
        return answer_text, f"?/{marks}"

    from shark_answer.config import ModelProvider
    scorer = _registry.get(ModelProvider.CLAUDE)
    if not scorer:
        return answer_text, f"?/{marks}"

    best_answer = answer_text
    best_achieved = 0
    score_str = f"?/{marks}"
    lang_note = " (Answer is in Chinese — evaluate accordingly.)" if language == "zh" else ""

    for round_num in range(3):
        score_prompt = (
            f"Question [{marks} marks]:\n{question_text}\n\n"
            f"Student answer:{lang_note}\n{best_answer}"
        )
        try:
            score_resp = await scorer.generate(
                prompt=score_prompt,
                system=SCORE_SYSTEM,
                temperature=0.1,
                max_tokens=512,
            )
            _cost_tracker.record(score_resp, subject, "scoring")
        except Exception:
            break

        if not score_resp.success:
            break

        try:
            raw = score_resp.content
            start = raw.find("{")
            end = raw.rfind("}") + 1
            if start < 0 or end <= start:
                break
            score_data = json.loads(raw[start:end])
            achieved = int(score_data.get("marks_achieved", 0))
            total    = int(score_data.get("total_marks", marks))
            verdict  = score_data.get("verdict", "partial")
            missing  = score_data.get("missing_points", [])

            if achieved > best_achieved:
                best_achieved = achieved
                score_str = f"{achieved}/{total}"

            if verdict == "full" or not missing or round_num == 2:
                break   # done — full marks or nothing left to improve

            # Revise
            lang_suffix = "\n\nWrite in Chinese (简体中文)." if language == "zh" else ""
            revise_prompt = (
                f"Question [{marks} marks]:\n{question_text}\n\n"
                f"Your current answer:\n{best_answer}\n\n"
                f"Examiner feedback — missing mark points:\n"
                + "\n".join(f"• {p}" for p in missing)
                + f"\n\nRewrite to earn these missing marks.{lang_suffix}"
            )
            try:
                revise_resp = await scorer.generate(
                    prompt=revise_prompt,
                    system=REVISE_SYSTEM,
                    temperature=0.3,
                    max_tokens=4096,
                )
                _cost_tracker.record(revise_resp, subject, "revision")
            except Exception:
                break

            if revise_resp.success and revise_resp.content.strip():
                best_answer = revise_resp.content

        except (json.JSONDecodeError, ValueError, KeyError):
            break

    return best_answer, score_str


# ===========================================================
# Pipeline dispatch helper (shared by /api/solve and /api/solve/stream)
# ===========================================================

async def _run_pipeline(
    q,
    pipeline,
    subject: str,
    lang,
    max_versions: int,
    profile,
    paper_number: Optional[int] = None,
) -> "PipelineResult":
    """Dispatch a single question to the appropriate pipeline.

    kb_context is fetched once here and passed into the pipeline so every
    model call in the pipeline automatically gets the relevant mark schemes,
    examiner reports, grade thresholds, and syllabus as context.
    """
    kb_context = build_prompt_context(subject=subject, paper_number=paper_number)

    if pipeline == Pipeline.SCIENCE_MATH:
        return await run_pipeline_a(
            question=q, subject=subject, registry=_registry,
            config=_config, cost_tracker=_cost_tracker,
            kb_context=kb_context,
            examiner_profile=profile,
            language=lang.value, max_versions=max_versions,
        )
    elif pipeline == Pipeline.ESSAY:
        return await run_pipeline_b(
            question=q, subject=subject, registry=_registry,
            config=_config, cost_tracker=_cost_tracker,
            kb_context=kb_context,
            examiner_profile=profile,
            language=lang.value, max_versions=max_versions,
        )
    elif pipeline == Pipeline.CS:
        return await run_pipeline_c(
            question=q, subject=subject, registry=_registry,
            config=_config, cost_tracker=_cost_tracker,
            kb_context=kb_context,
            examiner_profile=profile,
            language=lang.value, max_versions=max_versions,
        )
    else:
        from shark_answer.pipelines.base import PipelineResult as PR
        return PR(
            question=q, pipeline=pipeline.value, subject=subject,
            errors=[f"Pipeline {pipeline.value} not implemented"],
        )


def _build_question_results(results: list) -> list[QuestionResultOut]:
    """Convert PipelineResult list → QuestionResultOut list."""
    out = []
    for pr in results:
        versions_out = [
            AnswerVersionOut(
                version_number=v.version_number,
                answer_text=v.answer_text,
                explanation_text=v.explanation_text,
                approach_label=v.approach_label,
                provider=v.provider,
                verified=v.verified,
                quality_score=v.quality_score,
                language=v.language,
            )
            for v in pr.versions
        ]
        out.append(QuestionResultOut(
            question_number=pr.question.number,
            question_text=pr.question.text,
            pipeline=pr.pipeline,
            subject=pr.subject,
            versions=versions_out,
            verification_notes=pr.verification_notes,
            disagreement_resolved=pr.disagreement_resolved,
            errors=pr.errors,
        ))
    return out


# ===========================================================
# Frontend routes
# ===========================================================

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """Serve the main UI."""
    return templates.TemplateResponse(request, "index.html")


@app.get("/preview/{submission_id}", response_class=HTMLResponse)
async def print_preview(request: Request, submission_id: str):
    """Print-friendly preview of a submission."""
    entry = _find_history(submission_id)
    if not entry:
        raise HTTPException(404, "Submission not found")
    return templates.TemplateResponse(request, "preview.html", {"data": entry})


# ===========================================================
# API endpoints
# ===========================================================

@app.post("/api/solve", response_model=PaperResultOut)
async def solve_exam_paper(
    images: list[UploadFile] = File(..., description="Exam paper files (images/PDF/DOCX)"),
    subject: str = Form(..., description="Subject"),
    language: str = Form("en", description="Output language: en or zh"),
    examiner_profile: str = Form("", description="Examiner profile name (optional)"),
    max_versions: int = Form(5, description="Max answer versions per question (1-5)"),
):
    """Upload exam paper files and get A/A* answers.

    Accepted file types: JPG, PNG, WEBP, HEIC, PDF, DOCX.
    PDFs are split into one image per page; DOCX text and images are extracted.
    """
    global _last_run_log
    if not _config or not _registry or not _cost_tracker:
        raise HTTPException(500, "Server not initialized")

    try:
        subject_enum = Subject(subject)
    except ValueError:
        raise HTTPException(
            400,
            f"Invalid subject: {subject}. Valid: {[s.value for s in Subject]}",
        )

    lang = Language.ZH if language == "zh" else Language.EN
    max_versions = max(1, min(5, max_versions))
    _last_run_log = {"started": time.strftime("%Y-%m-%d %H:%M:%S"), "subject": subject, "stages": []}

    # Read and convert all uploaded files to image bytes
    image_data_list: list[bytes] = []
    filenames: list[str] = []
    for upload in images:
        raw_data = await upload.read()
        if not raw_data:
            raise HTTPException(400, f"Empty file: {upload.filename}")
        fname = upload.filename or "unknown"
        filenames.append(fname)
        converted = convert_file_to_images(fname, raw_data)
        if not converted:
            logger.warning("Could not convert file '%s'; skipping.", fname)
            continue
        image_data_list.extend(converted)

    if not image_data_list:
        raise HTTPException(422, "No usable image content could be extracted from the uploaded files.")

    # Step 1: Extract questions via AI vision
    logger.info("Extracting questions from %d image(s) for subject '%s'", len(image_data_list), subject)
    questions, extract_responses = await extract_questions_from_images(_registry, image_data_list)
    _cost_tracker.record_batch(extract_responses, subject, "extraction")
    _last_run_log["stages"].append({"stage": "extract", "questions": len(questions)})

    if not questions:
        raise HTTPException(422, "Could not extract any questions from the uploaded files")

    # Step 2: Load examiner profile
    profile = None
    if examiner_profile and _examiner_manager:
        profile = _examiner_manager.get_profile(examiner_profile)
        if not profile:
            profile = _examiner_manager.get_profile_for_subject(subject)

    # Step 3: Process each question through its pipeline
    results: list[PipelineResult] = []
    for q in questions:
        pipeline = route_question(q, subject_enum)
        logger.info("Processing Q%s via Pipeline %s", q.number, pipeline.value)
        try:
            pr = await _run_pipeline(q, pipeline, subject, lang, max_versions, profile)
            results.append(pr)
        except Exception as e:
            logger.exception("Error processing Q%s", q.number)
            results.append(PipelineResult(
                question=q, pipeline=pipeline.value, subject=subject, errors=[str(e)],
            ))

    responding = sum(1 for pr in results if pr.versions)
    _last_run_log["stages"].append({
        "stage": "solve", "questions": len(questions), "with_answers": responding,
    })

    # Step 4: Score and revise each answer version
    for pr in results:
        for v in pr.versions:
            try:
                new_text, score_str = await _score_and_revise(
                    question_text=pr.question.text,
                    marks=pr.question.marks,
                    answer_text=v.answer_text,
                    subject=subject,
                    language=lang.value,
                )
                v.answer_text = new_text
                v.quality_score = score_str
            except Exception as exc:
                logger.warning("Scoring failed for Q%s V%s: %s", pr.question.number, v.version_number, exc)

    _last_run_log["stages"].append({"stage": "score", "done": True})
    _last_run_log["completed"] = time.strftime("%Y-%m-%d %H:%M:%S")

    # Build response
    question_results = _build_question_results(results)
    paper_result = PaperResultOut(
        total_questions=len(questions),
        results=question_results,
        cost_summary=_cost_tracker.summary(),
    )

    # Save to history
    sid = str(uuid.uuid4())[:8]
    _history.append({
        "id": sid,
        "subject": subject,
        "language": language,
        "filenames": filenames,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "total_questions": paper_result.total_questions,
        "data": paper_result.model_dump(),
    })

    return paper_result


# ===========================================================
# SSE streaming solve endpoint
# ===========================================================

@app.post("/api/solve/stream")
async def solve_exam_paper_stream(
    images: list[UploadFile] = File(..., description="Exam paper files (images/PDF/DOCX)"),
    subject: str = Form(..., description="Subject"),
    language: str = Form("en", description="Output language: en or zh"),
    examiner_profile: str = Form("", description="Examiner profile name (optional)"),
    max_versions: int = Form(5, description="Max answer versions per question (1-5)"),
):
    """SSE streaming version of /api/solve — yields real-time progress events.

    Event types: stage, progress, scoring, scored, result, error
    """
    if not _config or not _registry or not _cost_tracker:
        raise HTTPException(500, "Server not initialized")

    try:
        subject_enum = Subject(subject)
    except ValueError:
        raise HTTPException(400, f"Invalid subject: {subject}. Valid: {[s.value for s in Subject]}")

    lang = Language.ZH if language == "zh" else Language.EN
    max_versions = max(1, min(5, max_versions))

    # Read all files upfront — must finish before StreamingResponse starts
    image_data_list: list[bytes] = []
    filenames: list[str] = []
    for upload in images:
        raw_data = await upload.read()
        if not raw_data:
            raise HTTPException(400, f"Empty file: {upload.filename}")
        fname = upload.filename or "unknown"
        filenames.append(fname)
        converted = convert_file_to_images(fname, raw_data)
        if converted:
            image_data_list.extend(converted)

    if not image_data_list:
        raise HTTPException(422, "No usable image content could be extracted.")

    profile_name = examiner_profile
    return StreamingResponse(
        _solve_sse_generator(
            image_data_list, filenames, subject_enum, lang,
            max_versions, profile_name, subject,
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


async def _solve_sse_generator(
    image_data_list: list[bytes],
    filenames: list[str],
    subject_enum,
    lang,
    max_versions: int,
    profile_name: str,
    subject: str,
) -> AsyncGenerator[str, None]:
    """Async generator that yields SSE-formatted strings during the solve pipeline."""
    global _last_run_log

    def sse(event_name: str, data: dict) -> str:
        return f"event: {event_name}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

    _last_run_log = {"started": time.strftime("%Y-%m-%d %H:%M:%S"), "subject": subject, "stages": []}

    try:
        # ── Stage: extract ──
        yield sse("stage", {"id": "extract", "label": "Extracting questions via AI vision...", "done": False})
        questions, extract_responses = await extract_questions_from_images(_registry, image_data_list)
        _cost_tracker.record_batch(extract_responses, subject, "extraction")
        _last_run_log["stages"].append({"stage": "extract", "questions": len(questions)})

        if not questions:
            yield sse("error", {"message": "Could not extract any questions from the uploaded files"})
            return

        yield sse("stage", {
            "id": "extract",
            "label": f"Extracted {len(questions)} question(s) ✓",
            "done": True,
        })

        # Load profile
        profile = None
        if profile_name and _examiner_manager:
            profile = _examiner_manager.get_profile(profile_name) or _examiner_manager.get_profile_for_subject(subject)

        # ── Stage: solve ──
        yield sse("stage", {
            "id": "solve",
            "label": f"Solving {len(questions)} question(s) across all models...",
            "done": False,
        })

        results: list[PipelineResult] = []
        for i, q in enumerate(questions):
            yield sse("progress", {
                "question": i + 1,
                "total": len(questions),
                "label": f"Solving Q{q.number} ({q.marks} marks)…",
            })
            pipeline = route_question(q, subject_enum)
            try:
                pr = await _run_pipeline(q, pipeline, subject, lang, max_versions, profile)
            except Exception as e:
                logger.exception("Error processing Q%s", q.number)
                pr = PipelineResult(
                    question=q, pipeline=pipeline.value, subject=subject, errors=[str(e)],
                )
            results.append(pr)

        responding = sum(1 for pr in results if pr.versions)
        _last_run_log["stages"].append({
            "stage": "solve", "questions": len(questions), "with_answers": responding,
        })
        yield sse("stage", {
            "id": "solve",
            "label": f"Solved {len(questions)} question(s) — models responding: {responding}/{len(questions)} ✓",
            "done": True,
        })

        # ── Stage: score ──
        total_versions = sum(len(pr.versions) for pr in results)
        yield sse("stage", {
            "id": "score",
            "label": f"Scoring {total_versions} answer version(s)…",
            "done": False,
        })

        scored = 0
        for pr in results:
            for v in pr.versions:
                yield sse("scoring", {
                    "question": pr.question.number,
                    "version": v.version_number,
                    "label": f"Scoring Q{pr.question.number} V{v.version_number}…",
                })
                try:
                    new_text, score_str = await _score_and_revise(
                        question_text=pr.question.text,
                        marks=pr.question.marks,
                        answer_text=v.answer_text,
                        subject=subject,
                        language=lang.value,
                    )
                    v.answer_text = new_text
                    v.quality_score = score_str
                except Exception as exc:
                    logger.warning("Scoring Q%s V%s: %s", pr.question.number, v.version_number, exc)
                    score_str = f"?/{pr.question.marks}"
                    v.quality_score = score_str
                scored += 1
                yield sse("scored", {
                    "question": pr.question.number,
                    "version": v.version_number,
                    "score": score_str,
                })

        _last_run_log["stages"].append({"stage": "score", "versions_scored": scored})
        yield sse("stage", {"id": "score", "label": f"Scoring complete — {scored} version(s) scored ✓", "done": True})

        # ── Stage: finalize ──
        yield sse("stage", {"id": "finalize", "label": "Building final results…", "done": False})

        question_results = _build_question_results(results)
        paper_result = PaperResultOut(
            total_questions=len(questions),
            results=question_results,
            cost_summary=_cost_tracker.summary(),
        )

        sid = str(uuid.uuid4())[:8]
        _history.append({
            "id": sid,
            "subject": subject,
            "language": lang.value,
            "filenames": filenames,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "total_questions": paper_result.total_questions,
            "data": paper_result.model_dump(),
        })

        _last_run_log["completed"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _last_run_log["submission_id"] = sid

        yield sse("stage", {"id": "finalize", "label": "Complete ✓", "done": True})

        result_data = paper_result.model_dump()
        result_data["submission_id"] = sid
        yield sse("result", result_data)

    except Exception as e:
        logger.exception("SSE solve failed")
        yield sse("error", {"message": str(e)})


@app.get("/api/history")
async def get_history():
    """Return submission history (without full data payloads)."""
    return [
        {
            "id": h["id"],
            "subject": h["subject"],
            "language": h["language"],
            "filenames": h["filenames"],
            "timestamp": h["timestamp"],
            "total_questions": h["total_questions"],
        }
        for h in reversed(_history)
    ]


@app.get("/api/history/{submission_id}")
async def get_history_entry(submission_id: str):
    """Return a specific submission's full data."""
    entry = _find_history(submission_id)
    if not entry:
        raise HTTPException(404, "Submission not found")
    return entry["data"]


# ===========================================================
# Chat / Answer Correction endpoint
# ===========================================================

@app.post("/api/chat", response_model=ChatResponse)
async def chat_correction(req: ChatRequest):
    """Answer correction chat powered by Claude as judge.

    Send a question + user message to Claude; it reviews the existing
    answer and returns a correction or explanation.
    Conversation history is kept per (submission_id, question_number).
    """
    # Check submission first — this 404 can happen without AI config
    entry = _find_history(req.submission_id)
    if not entry:
        raise HTTPException(404, "Submission not found")

    if not _config or not _registry or not _cost_tracker:
        raise HTTPException(500, "Server not initialized")

    # Find the specific question result
    qr = next(
        (r for r in entry["data"]["results"]
         if r["question_number"] == req.question_number),
        None,
    )
    if not qr:
        raise HTTPException(404, f"Question {req.question_number} not found in submission")

    # Build chat key
    chat_key = f"{req.submission_id}:{req.question_number}"
    history = _chat_histories.setdefault(chat_key, [])

    # Build context: first version answer (best available)
    best_version = qr["versions"][0] if qr["versions"] else {}
    answer_context = best_version.get("answer_text", "(no answer)")

    # Build the judge prompt
    system_prompt = (
        "You are an expert A-Level CIE examiner and academic tutor. "
        "A student is asking you to review or correct an AI-generated exam answer. "
        "Be precise, cite marking criteria where relevant, and offer concise corrections. "
        "If the answer is correct, confirm it and explain why marks would be awarded. "
        "If the answer is wrong or incomplete, give the correct answer with reasoning. "
        "Keep responses focused and exam-appropriate."
    )

    # Build conversation messages
    messages_for_api: list[dict] = []

    # First message includes full context
    if not history:
        initial_ctx = (
            f"Question {req.question_number}: {qr['question_text']}\n\n"
            f"Subject: {qr['subject']} | Pipeline: {qr['pipeline']}\n\n"
            f"Generated Answer (Version 1):\n{answer_context}"
        )
        messages_for_api.append({
            "role": "user",
            "content": f"[Context]\n{initial_ctx}\n\n[Student question]\n{req.message}",
        })
    else:
        # Replay previous turns
        for turn in history:
            messages_for_api.append({
                "role": turn["role"],
                "content": turn["content"],
            })
        messages_for_api.append({"role": "user", "content": req.message})

    # Call Claude as judge via the provider registry
    try:
        from shark_answer.config import ModelProvider

        claude = _registry.get(ModelProvider.CLAUDE)
        if claude is None:
            raise HTTPException(503, "Claude provider not configured")

        # Build a single prompt string (base generate() doesn't support multi-turn messages)
        if not history:
            prompt = (
                f"Context:\n"
                f"Question {req.question_number}: {qr['question_text']}\n"
                f"Subject: {qr['subject']} | Pipeline: {qr['pipeline']}\n\n"
                f"Generated Answer (Version 1):\n{answer_context}\n\n"
                f"Student question: {req.message}"
            )
        else:
            parts: list[str] = []
            for i, turn in enumerate(history):
                content = turn["content"]
                if i == 0 and content.startswith("[Context]"):
                    p = content.split("[Student question]\n", 1)
                    content = p[1] if len(p) > 1 else content
                prefix = "Student" if turn["role"] == "user" else "You"
                parts.append(f"{prefix}: {content}")
            parts.append(f"Student: {req.message}")
            prompt = "\n\n".join(parts)

        response = await claude.generate(
            prompt=prompt,
            system=system_prompt,
            temperature=0.3,
            max_tokens=2048,
        )

        if not response.success:
            raise HTTPException(503, f"Judge model error: {response.error}")

        reply_text = response.content
        _cost_tracker.record(response, qr["subject"], "chat_correction")

    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Chat correction failed")
        raise HTTPException(500, f"Chat error: {exc}") from exc

    # Persist history
    if not history:
        history.append({
            "role": "user",
            "content": f"[Context]\nQuestion {req.question_number}: {qr['question_text']}\n\n[Student question]\n{req.message}",
        })
    else:
        history.append({"role": "user", "content": req.message})
    history.append({"role": "assistant", "content": reply_text})

    # Return public-facing history (strip context prefix from first msg)
    public_history = []
    for i, turn in enumerate(history):
        content = turn["content"]
        if i == 0 and turn["role"] == "user" and content.startswith("[Context]"):
            parts_pub = content.split("[Student question]\n", 1)
            content = parts_pub[1] if len(parts_pub) > 1 else content
        public_history.append({"role": turn["role"], "content": content})

    return ChatResponse(reply=reply_text, chat_history=public_history)


@app.get("/api/chat/{submission_id}/{question_number}")
async def get_chat_history(submission_id: str, question_number: str):
    """Retrieve existing chat history for a question."""
    chat_key = f"{submission_id}:{question_number}"
    raw = _chat_histories.get(chat_key, [])
    # Strip system context from first user message for public display
    public = []
    for i, turn in enumerate(raw):
        content = turn["content"]
        if i == 0 and turn["role"] == "user" and content.startswith("[Context]"):
            parts = content.split("[Student question]\n", 1)
            content = parts[1] if len(parts) > 1 else content
        public.append({"role": turn["role"], "content": content})
    return public


# ===========================================================
# Export endpoints
# ===========================================================

@app.post("/api/export/pdf")
async def export_pdf(request: Request):
    """Export results as a formatted PDF."""
    body = await request.json()
    data = body.get("data")
    if not data:
        raise HTTPException(400, "Missing data")
    pdf_bytes = _generate_pdf(data)
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=shark_answer_results.pdf"},
    )


@app.post("/api/export/docx")
async def export_docx(request: Request):
    """Export results as a Word document."""
    body = await request.json()
    data = body.get("data")
    if not data:
        raise HTTPException(400, "Missing data")
    docx_bytes = _generate_docx(data)
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=shark_answer_results.docx"},
    )


@app.post("/api/export/md")
async def export_markdown(request: Request):
    """Export results as Markdown (.md)."""
    body = await request.json()
    data = body.get("data")
    if not data:
        raise HTTPException(400, "Missing data")
    md_text = _generate_markdown(data)
    return StreamingResponse(
        io.BytesIO(md_text.encode("utf-8")),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=shark_answer_results.md"},
    )


@app.post("/api/export/txt")
async def export_txt(request: Request):
    """Export results as plain text (.txt)."""
    body = await request.json()
    data = body.get("data")
    if not data:
        raise HTTPException(400, "Missing data")
    txt = _generate_txt(data)
    return StreamingResponse(
        io.BytesIO(txt.encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=shark_answer_results.txt"},
    )


@app.post("/api/export/xlsx")
async def export_xlsx(request: Request):
    """Export results as an Excel spreadsheet (.xlsx)."""
    body = await request.json()
    data = body.get("data")
    if not data:
        raise HTTPException(400, "Missing data")
    xlsx_bytes = _generate_xlsx(data)
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=shark_answer_results.xlsx"},
    )


# ===========================================================
# Remaining API endpoints
# ===========================================================

@app.post("/api/export/pdf/paper")
async def export_pdf_paper(request: Request):
    """Export a single paper version as a clean student answer booklet PDF.

    Body: {data: PaperResultOut, version_number: int}
    """
    body = await request.json()
    data = body.get("data")
    version_number = int(body.get("version_number", 1))
    if not data:
        raise HTTPException(400, "Missing data")
    subject = (data.get("results", [{}])[0].get("subject", "paper") if data.get("results") else "paper")
    pdf_bytes = _generate_paper_pdf(data, version_number, subject)
    filename = f"paper_{version_number}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/api/export/pdf/zip")
async def export_pdf_zip(request: Request):
    """Export multiple paper versions as a ZIP of individual PDFs.

    Body: {data: PaperResultOut, version_numbers: [1, 2, 3]}
    """
    body = await request.json()
    data = body.get("data")
    version_numbers: list[int] = [int(v) for v in body.get("version_numbers", [1])]
    if not data:
        raise HTTPException(400, "Missing data")
    subject = (data.get("results", [{}])[0].get("subject", "paper") if data.get("results") else "paper")

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for vn in version_numbers:
            pdf_bytes = _generate_paper_pdf(data, vn, subject)
            zf.writestr(f"paper_{vn}.pdf", pdf_bytes)
    zip_buf.seek(0)
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=shark_papers.zip"},
    )


@app.post("/api/practical/predict")
async def predict_practical(
    image: Optional[UploadFile] = File(None),
    text_content: str = Form(""),
):
    """Pipeline D: Predict physics practical experiment from equipment list."""
    if not _config or not _registry or not _cost_tracker:
        raise HTTPException(500, "Server not initialized")
    image_data = None
    if image:
        image_data = await image.read()
    if not image_data and not text_content:
        raise HTTPException(400, "Provide either an image or text content")

    result = await run_pipeline_d(
        image_data=image_data,
        text_content=text_content or None,
        registry=_registry,
        config=_config,
        cost_tracker=_cost_tracker,
    )
    return {
        "equipment_detected": result.equipment_detected,
        "predictions": [asdict(p) for p in result.predictions],
        "errors": result.errors,
        "cost_summary": _cost_tracker.summary(),
    }


@app.get("/api/cost", response_model=CostSummaryOut)
async def get_cost_summary():
    """Get current session cost breakdown."""
    if not _cost_tracker:
        raise HTTPException(500, "Server not initialized")
    data = _cost_tracker.summary()
    return CostSummaryOut(**data)


@app.get("/api/examiner/profiles", response_model=list[ExaminerProfileOut])
async def list_examiner_profiles():
    """List all examiner profiles."""
    if not _examiner_manager:
        raise HTTPException(500, "Server not initialized")
    return [
        ExaminerProfileOut(
            name=p.name, subject=p.subject,
            region=p.region, description=p.description,
        )
        for p in _examiner_manager.list_profiles()
    ]


@app.post("/api/examiner/profiles")
async def create_examiner_profile(
    name: str = Form(...),
    subject: str = Form(...),
    region: str = Form("default"),
    description: str = Form(""),
    evaluation_depth: float = Form(0.5),
    real_world_examples: float = Form(0.5),
    diagram_preference: float = Form(0.5),
    formula_rigour: float = Form(0.5),
    penalizes_formulaic: bool = Form(False),
    values_originality: bool = Form(False),
    strict_on_units: bool = Form(True),
    custom_instructions: str = Form(""),
):
    """Create a new examiner profile."""
    if not _examiner_manager:
        raise HTTPException(500, "Server not initialized")
    from shark_answer.modules.examiner_profile import ExaminerProfile

    profile = ExaminerProfile(
        name=name, subject=subject, region=region,
        description=description, evaluation_depth=evaluation_depth,
        real_world_examples=real_world_examples,
        diagram_preference=diagram_preference,
        formula_rigour=formula_rigour,
        penalizes_formulaic=penalizes_formulaic,
        values_originality=values_originality,
        strict_on_units=strict_on_units,
        custom_instructions=custom_instructions,
    )
    _examiner_manager.add_profile(profile)
    return {"status": "created", "name": name}


@app.get("/api/models")
async def list_available_models():
    """List configured and available models."""
    from shark_answer.config import PIPELINE_MODEL_CONFIG

    if not _config:
        raise HTTPException(500, "Server not initialized")
    return {
        "configured": [m.value for m in _config.models.keys()],
        "pipeline_assignments": {
            p.value: {
                tier: [m.value for m in _config.get_pipeline_models(p, tier)]
                for tier in conf.keys()
            }
            for p, conf in PIPELINE_MODEL_CONFIG.items()
        },
    }


@app.get("/api/health")
async def health():
    """Health check."""
    configured = len(_config.models) if _config else 0
    return {
        "status": "ok",
        "models_configured": configured,
        "total_cost": _cost_tracker.total_cost if _cost_tracker else 0,
    }


@app.get("/api/debug/last-run")
async def debug_last_run():
    """Return a log of the most recent /api/solve or /api/solve/stream run."""
    return _last_run_log


@app.get("/api/debug/models")
async def debug_models():
    """Test every configured model with a hello-world prompt.

    Returns pass/fail/error for each model so you can see which providers
    are actually working vs silently failing.
    """
    if not _config or not _registry:
        raise HTTPException(500, "Server not initialized")

    results: dict[str, dict] = {}
    for provider in list(_config.models.keys()):
        inst = _registry.get(provider)
        if inst is None:
            results[provider.value] = {"status": "not_configured", "model": ""}
            continue
        try:
            resp = await inst.generate(
                prompt="Hello. Please respond with just the single word: OK",
                system="",
                temperature=0.0,
                max_tokens=20,
            )
            results[provider.value] = {
                "status": "ok" if resp.success else "error",
                "model": resp.model_name,
                "response": resp.content[:120].strip() if resp.success else None,
                "error": resp.error if not resp.success else None,
                "latency_s": round(resp.latency_seconds, 2),
                "input_tokens": resp.usage.input_tokens,
                "output_tokens": resp.usage.output_tokens,
            }
        except Exception as exc:
            results[provider.value] = {
                "status": "exception",
                "model": getattr(inst, "model_name", ""),
                "error": str(exc),
            }

    ok_count = sum(1 for v in results.values() if v["status"] == "ok")
    return {
        "total_configured": len(_config.models),
        "total_ok": ok_count,
        "models": results,
    }


# ===========================================================
# Helper functions
# ===========================================================

def _find_history(sid: str) -> Optional[dict]:
    for h in _history:
        if h["id"] == sid:
            return h
    return None


def _strip_md(text: str) -> str:
    """Strip basic Markdown formatting for plain text export."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"#{1,6}\s+", "", text)
    text = re.sub(r"`{1,3}[^`]*`{1,3}", lambda m: m.group(0).strip("`"), text)
    return text


# ---- Markdown export ----

def _generate_markdown(data: dict) -> str:
    lines: list[str] = []
    subject = (
        data.get("results", [{}])[0].get("subject", "Unknown")
        if data.get("results") else "Unknown"
    )
    lines.append("# Shark Answer — Exam Results\n")
    lines.append(
        f"**Subject:** {subject.replace('_', ' ').title()}  "
        f"**Questions:** {data.get('total_questions', 0)}  "
        f"**Cost:** ${data.get('cost_summary', {}).get('total_cost_usd', 0):.4f}\n"
    )
    lines.append("---\n")

    for qr in data.get("results", []):
        lines.append(f"## Question {qr['question_number']}\n")
        lines.append(f"{qr.get('question_text', '')}\n")
        if qr.get("errors"):
            lines.append(f"> ⚠️ Errors: {'; '.join(qr['errors'])}\n")

        for v in qr.get("versions", []):
            badge = " ✓" if v.get("verified") else ""
            score = f" · Score: {v['quality_score']}" if v.get("quality_score") is not None else ""
            lines.append(
                f"### Version {v['version_number']}: {v.get('approach_label', '')}"
                f"{badge}{score}\n"
            )
            lines.append(f"*Provider: {v.get('provider', '')}*\n")
            lines.append(f"{v.get('answer_text', '')}\n")
            if v.get("explanation_text"):
                lines.append(f"\n**Explanation:**\n{v['explanation_text']}\n")
            lines.append("\n---\n")

    return "\n".join(lines)


# ---- Plain text export ----

def _generate_txt(data: dict) -> str:
    lines: list[str] = []
    subject = (
        data.get("results", [{}])[0].get("subject", "Unknown")
        if data.get("results") else "Unknown"
    )
    sep = "=" * 60
    thin = "-" * 40

    lines.append("SHARK ANSWER — EXAM RESULTS")
    lines.append(sep)
    lines.append(
        f"Subject: {subject.replace('_', ' ').title()}  "
        f"Questions: {data.get('total_questions', 0)}  "
        f"Cost: ${data.get('cost_summary', {}).get('total_cost_usd', 0):.4f}"
    )
    lines.append(sep)
    lines.append("")

    for qr in data.get("results", []):
        lines.append(f"QUESTION {qr['question_number']}")
        lines.append(thin)
        lines.append(_strip_md(qr.get("question_text", "")))
        lines.append("")

        for v in qr.get("versions", []):
            badge = " [VERIFIED]" if v.get("verified") else ""
            score = f" Score:{v['quality_score']}" if v.get("quality_score") is not None else ""
            lines.append(
                f"  Version {v['version_number']}: "
                f"{v.get('approach_label', '')} ({v.get('provider', '')}){badge}{score}"
            )
            lines.append("")
            for para in _strip_md(v.get("answer_text", "")).split("\n"):
                lines.append(f"  {para}")
            lines.append("")
            if v.get("explanation_text"):
                lines.append("  [Explanation]")
                for para in _strip_md(v["explanation_text"]).split("\n"):
                    lines.append(f"  {para}")
                lines.append("")

        lines.append(sep)
        lines.append("")

    return "\n".join(lines)


# ---- Excel export ----

def _generate_xlsx(data: dict) -> bytes:
    """Generate an Excel workbook with one row per answer version."""
    from openpyxl import Workbook  # type: ignore
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side  # type: ignore

    wb = Workbook()
    ws = wb.active
    ws.title = "Results"

    # Header style
    hdr_fill = PatternFill("solid", fgColor="1E3A8A")  # dark blue
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers = [
        "Question #", "Question Text", "Pipeline", "Subject",
        "Version", "Approach", "Provider", "Verified",
        "Quality Score", "Language", "Answer", "Explanation",
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_border

    # Set column widths
    col_widths = [10, 40, 12, 16, 8, 20, 16, 10, 12, 10, 60, 60]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w
    ws.row_dimensions[1].height = 28

    # Data rows
    row_fill_a = PatternFill("solid", fgColor="F0F4FF")
    row_fill_b = PatternFill("solid", fgColor="FFFFFF")
    data_align = Alignment(vertical="top", wrap_text=True)

    row_num = 2
    for qi, qr in enumerate(data.get("results", [])):
        for v in qr.get("versions", []):
            fill = row_fill_a if qi % 2 == 0 else row_fill_b
            row_data = [
                qr["question_number"],
                qr.get("question_text", ""),
                qr.get("pipeline", ""),
                qr.get("subject", ""),
                v.get("version_number", ""),
                v.get("approach_label", ""),
                v.get("provider", ""),
                "Yes" if v.get("verified") else "No",
                v.get("quality_score", ""),
                v.get("language", ""),
                v.get("answer_text", ""),
                v.get("explanation_text", ""),
            ]
            ws.append(row_data)
            for cell in ws[row_num]:
                cell.fill = fill
                cell.alignment = data_align
                cell.border = thin_border
            ws.row_dimensions[row_num].height = 60
            row_num += 1

    # Summary sheet
    ws2 = wb.create_sheet("Summary")
    cost = data.get("cost_summary", {})
    ws2.append(["Metric", "Value"])
    for cell in ws2[1]:
        cell.font = Font(bold=True)
    ws2.append(["Total Questions", data.get("total_questions", 0)])
    ws2.append(["Total Cost (USD)", f"${cost.get('total_cost_usd', 0):.4f}"])
    ws2.append(["Total API Calls", cost.get("total_calls", 0)])
    ws2.append(["Total Input Tokens", cost.get("total_input_tokens", 0)])
    ws2.append(["Total Output Tokens", cost.get("total_output_tokens", 0)])
    ws2.column_dimensions["A"].width = 24
    ws2.column_dimensions["B"].width = 20

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---- PDF helpers ----

def _strip_latex_for_pdf(text: str) -> str:
    """Convert LaTeX delimiters to plain text for PDF rendering.

    Strips $ and $$ markers but keeps the formula content readable.
    E.g.  $F = ma$  →  F = ma
          $$\\int_0^1 x dx$$  →  ∫₀¹ x dx (simplified)
    """
    import re
    # Display math $$...$$
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: f'[{m.group(1).strip()}]', text, flags=re.DOTALL)
    # Inline math $...$
    text = re.sub(r'\$(.+?)\$', lambda m: m.group(1).strip(), text)
    # Remove \left, \right, \mathrm{} etc — keep content
    text = re.sub(r'\\(?:left|right|mathrm|mathbf|text)\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\(?:left|right)[.()\[\]|]', '', text)
    # Replace common LaTeX macros
    text = text.replace('\\times', '×').replace('\\cdot', '·')
    text = text.replace('\\alpha', 'α').replace('\\beta', 'β').replace('\\gamma', 'γ')
    text = text.replace('\\delta', 'δ').replace('\\Delta', 'Δ').replace('\\pi', 'π')
    text = text.replace('\\mu', 'μ').replace('\\sigma', 'σ').replace('\\omega', 'ω')
    text = text.replace('\\Omega', 'Ω').replace('\\lambda', 'λ').replace('\\theta', 'θ')
    text = text.replace('\\int', '∫').replace('\\sum', '∑').replace('\\infty', '∞')
    text = text.replace('\\geq', '≥').replace('\\leq', '≤').replace('\\neq', '≠')
    text = text.replace('\\approx', '≈').replace('\\pm', '±').replace('\\sqrt', '√')
    text = text.replace('\\frac', '/').replace('^', '^').replace('_{', '_')
    # Strip remaining backslash commands
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)
    text = re.sub(r'[{}]', '', text)
    return text


def _pdf_safe(text: str) -> str:
    """Make text safe for reportlab (XML-escape + strip control chars)."""
    import re
    text = _strip_latex_for_pdf(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # Strip non-printable control characters (keep newlines/tabs)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text


def _try_register_unicode_font() -> str:
    """Try to register a Unicode-capable TTF font; return font name or 'Helvetica'."""
    import os
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        # macOS
        '/System/Library/Fonts/Supplemental/Arial Unicode MS.ttf',
        '/Library/Fonts/Arial Unicode MS.ttf',
        '/System/Library/Fonts/Times New Roman.ttf',
        '/System/Library/Fonts/Helvetica.ttc',
        # Linux
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf',
        # Windows
        'C:/Windows/Fonts/arial.ttf',
        'C:/Windows/Fonts/calibri.ttf',
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('UniFont', path))
                return 'UniFont'
            except Exception:
                continue
    return 'Helvetica'


# ---- Per-paper student answer booklet PDF ----

def _generate_paper_pdf(data: dict, version_number: int, subject: str) -> bytes:
    """Generate a clean student answer booklet PDF for a single paper version.

    No model names, no version style labels — just a clean answer sheet
    with subject header, paper number, sequential questions, and page numbers.
    """
    import datetime
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    body_font = _try_register_unicode_font()

    total_versions = max(
        (len(qr.get("versions", [])) for qr in data.get("results", [])),
        default=1,
    )
    date_str = datetime.date.today().strftime("%B %Y")
    subject_display = subject.replace("_", " ").title()

    buf = io.BytesIO()

    def _add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillGray(0.55)
        canvas.drawCentredString(A4[0] / 2, 10 * mm, str(doc.page))
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=25 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=22 * mm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "PH1", parent=styles["Normal"],
        fontSize=15, fontName=body_font,
        spaceAfter=2, alignment=TA_CENTER,
        textColor=colors.black,
    ))
    styles.add(ParagraphStyle(
        "PH2", parent=styles["Normal"],
        fontSize=9, fontName=body_font,
        spaceAfter=8, alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
    ))
    styles.add(ParagraphStyle(
        "QL", parent=styles["Normal"],
        fontSize=11, fontName=body_font,
        spaceBefore=14, spaceAfter=3,
        textColor=colors.black,
    ))
    styles.add(ParagraphStyle(
        "QT", parent=styles["Normal"],
        fontSize=9, fontName=body_font,
        spaceAfter=5, leftIndent=10,
        textColor=colors.HexColor("#555555"),
    ))
    styles.add(ParagraphStyle(
        "AB", parent=styles["Normal"],
        fontSize=11, leading=18, fontName=body_font,
        spaceAfter=8,
        textColor=colors.black,
    ))

    story: list = []
    story.append(Paragraph(_pdf_safe(subject_display), styles["PH1"]))
    story.append(Paragraph(
        f"Paper {version_number} of {total_versions} &nbsp;&middot;&nbsp; {date_str}",
        styles["PH2"],
    ))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
    story.append(Spacer(1, 6 * mm))

    for qr in data.get("results", []):
        versions = qr.get("versions", [])
        if not versions:
            continue
        # Pick the requested version index; fallback to last available
        v = (
            versions[version_number - 1]
            if version_number - 1 < len(versions)
            else versions[-1]
        )
        q_num = qr.get("question_number", "?")
        q_text = qr.get("question_text", "")

        story.append(Paragraph(_pdf_safe(f"Question {q_num}"), styles["QL"]))
        if q_text:
            story.append(Paragraph(_pdf_safe(q_text).replace("\n", "<br/>"), styles["QT"]))
        story.append(Spacer(1, 2 * mm))

        answer_text = v.get("answer_text", "(No answer generated)")
        story.append(Paragraph(_pdf_safe(answer_text).replace("\n", "<br/>"), styles["AB"]))
        story.append(Spacer(1, 4 * mm))

    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    return buf.getvalue()


# ---- PDF export ----

def _generate_pdf(data: dict) -> bytes:
    """Generate a PDF from results data using reportlab with Unicode font support."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    body_font = _try_register_unicode_font()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "SharkTitle", parent=styles["Heading1"],
        fontSize=20, spaceAfter=10, textColor=colors.HexColor("#1d4ed8"),
        fontName=body_font,
    ))
    styles.add(ParagraphStyle(
        "SharkH2", parent=styles["Heading2"],
        fontSize=13, spaceAfter=6, textColor=colors.HexColor("#1e40af"),
        fontName=body_font,
    ))
    styles.add(ParagraphStyle(
        "SharkH3", parent=styles["Heading3"],
        fontSize=11, spaceAfter=4, textColor=colors.HexColor("#2563eb"),
        fontName=body_font,
    ))
    styles.add(ParagraphStyle(
        "SharkBody", parent=styles["BodyText"],
        fontSize=10, leading=15, spaceAfter=5, fontName=body_font,
    ))
    styles.add(ParagraphStyle(
        "SharkMeta", parent=styles["BodyText"],
        fontSize=8, textColor=colors.grey, spaceAfter=4, fontName=body_font,
    ))

    story: list = []

    subject = (
        data.get("results", [{}])[0].get("subject", "Unknown")
        if data.get("results") else "Unknown"
    )
    story.append(Paragraph("Shark Answer — Exam Results", styles["SharkTitle"]))
    story.append(Paragraph(
        f"Subject: {subject.replace('_', ' ').title()} &nbsp;|&nbsp; "
        f"Questions: {data.get('total_questions', 0)} &nbsp;|&nbsp; "
        f"Cost: ${data.get('cost_summary', {}).get('total_cost_usd', 0):.4f}",
        styles["SharkMeta"],
    ))
    story.append(Spacer(1, 8))

    for qr in data.get("results", []):
        story.append(Paragraph(
            _pdf_safe(f"Question {qr['question_number']}"),
            styles["SharkH2"],
        ))
        story.append(Paragraph(
            _pdf_safe(qr.get("question_text", "")).replace("\n", "<br/>"),
            styles["SharkBody"],
        ))
        story.append(Spacer(1, 4))

        for v in qr.get("versions", []):
            badge = " [VERIFIED]" if v.get("verified") else ""
            lbl = _pdf_safe(
                f"Version {v['version_number']}: {v.get('approach_label', '')} "
                f"({v.get('provider', '')}){badge}"
            )
            story.append(Paragraph(f"<b>{lbl}</b>", styles["SharkH3"]))
            answer_text = _pdf_safe(v.get("answer_text", "")).replace("\n", "<br/>")
            story.append(Paragraph(answer_text, styles["SharkBody"]))
            story.append(Spacer(1, 4))

            if v.get("explanation_text"):
                story.append(Paragraph("<b>Explanation:</b>", styles["SharkBody"]))
                expl = _pdf_safe(v["explanation_text"]).replace("\n", "<br/>")
                story.append(Paragraph(expl, styles["SharkBody"]))
                story.append(Spacer(1, 6))

        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


# ---- DOCX export (unchanged) ----

def _generate_docx(data: dict) -> bytes:
    """Generate a Word document from results data."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("Shark Answer — Exam Results", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subject = (
        data.get("results", [{}])[0].get("subject", "Unknown")
        if data.get("results") else "Unknown"
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Subject: {subject.replace('_', ' ').title()}  |  "
        f"Questions: {data.get('total_questions', 0)}  |  "
        f"Cost: ${data.get('cost_summary', {}).get('total_cost_usd', 0):.4f}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    for qr in data.get("results", []):
        doc.add_heading(f"Question {qr['question_number']}", level=1)
        doc.add_paragraph(qr.get("question_text", ""))

        for v in qr.get("versions", []):
            badge = " [VERIFIED]" if v.get("verified") else ""
            heading = (
                f"Version {v['version_number']}: "
                f"{v.get('approach_label', '')} ({v.get('provider', '')}){badge}"
            )
            doc.add_heading(heading, level=2)
            doc.add_paragraph(v.get("answer_text", ""))

            if v.get("explanation_text"):
                p = doc.add_paragraph()
                run2 = p.add_run("Explanation")
                run2.bold = True
                run2.font.size = Pt(11)
                doc.add_paragraph(v["explanation_text"])

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
